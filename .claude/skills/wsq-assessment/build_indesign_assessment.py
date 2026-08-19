#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for
   'Creating Stunning Print and Digital Publications with InDesign' (TGS-2021007827).

The instruments MIRROR the paper the ATO has on file (confirmed against the LMS-TMS
record and Assessment Plan v5, 17 Oct 2025):

  · Practical Performance (PP) — 3 tasks, 75 minutes, covering abilities A1–A4
      Task 1 → (A1, A2)   Task 2 → (A3)   Task 3 → (A4)
  · Oral Questioning (OQ)      — 5 questions, 15 minutes, covering knowledge K1–K3
      Q1 → (K1)  Q2 → (K1)  Q3 → (K2)  Q4 → (K2)  Q5 → (K3)

There is NO written short-answer paper on this course. Instrument type, question
count, code mapping and timings are unchanged from the original; only the CONTENT is
rewritten against the current slides and Learner Guide activities.

Each instrument produces a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page. Page 1 cover; page 2 Trainee Information +
Instructions + Grading; the scenario/questions begin on page 3. Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
for _cand in (os.path.join(REPO, ".claude/skills/courseware-build/build"),
              os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc

# ─── COURSE ─────────────────────────────────────────────────────────────────
TITLE       = "Creating Stunning Print and Digital Publications with InDesign"
COURSE_CODE = "TGS-2021007827"
PP_MINUTES  = "75 minutes"
OQ_MINUTES  = "15 minutes"
# ────────────────────────────────────────────────────────────────────────────
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT = os.path.join(REPO, "assessment")
os.makedirs(OUT, exist_ok=True)

def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name),
              os.path.join(REPO, ".claude/skills/courseware-build/assets", name),
              os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = None

Q_VER, A_VER = "v5", "v5"
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)

# ---------------------------------------------------------------- ORAL QUESTIONING (KNOWLEDGE)
# 5 questions, mapped K1, K1, K2, K2, K3 — the same count and mapping as the original.
# Every answer is taught in the slides and the Learner Guide; the source is cited in the key.
ORAL = [
    ("K1",
     "InDesign carries forward the conventions of manual production drawing — a page is still "
     "planned before it is drawn.",
     "Before you create any InDesign document, what must you establish about the finished printed "
     "page, and why does each item matter to the printer?",
     ["Trim size — the finished cut size of the page. It is what you type as the page size in "
      "File > New > Document.",
      "Bleed — artwork extended past the trim, 3 mm in Singapore/ISO practice, so a small cutting "
      "drift never leaves a white sliver at the edge.",
      "Margin (safety) — the zone inside the trim that no live matter may cross, so text, logos and "
      "prices are never cut off.",
      "Slug — the area outside the bleed carrying the job name, colour bars and printer's marks; it "
      "is discarded after trimming.",
      "Binding method and page count, because they affect creep, gutter allowance and imposition.",
      "These come from the printer's written specification, not from assumption — a wrong trim or a "
      "missing bleed is only discovered at the press, after plates have been paid for.",
      "Taught in Topic 1 (Anatomy of a Print Page) and practised in Activity 1."]),

    ("K1",
     "A parent page (formerly called a master page) is the digital equivalent of the pre-ruled "
     "sheet a draughtsman worked on.",
     "What kinds of element belong on a parent page rather than on an individual document page, "
     "and what is the practical benefit of putting them there?",
     ["Repeating elements: the company logo, running heads and footers, the folio (page number), "
      "column and margin structure, and any rule or background that repeats across the publication.",
      "The page-number marker specifically — inserted with Type > Insert Special Character > "
      "Markers > Current Page Number, never typed as digits.",
      "Benefit: change the parent once and every page based on it updates automatically, so a "
      "200-page document is re-styled in seconds instead of page by page.",
      "It also guarantees consistency — the folio sits in exactly the same position on every page.",
      "Parent items appear on document pages with a dotted border and are locked; "
      "Ctrl/Cmd+Shift+click releases a single item where an override is genuinely needed.",
      "Taught in Topic 1 (Parent Pages, Automatic Page Numbering) and practised in Activity 3."]),

    ("K2",
     "InDesign is computer-aided drawing software, and threading is one of the processes that has "
     "no manual equivalent.",
     "What is threading, how do you create a thread between two text frames, and what does a red "
     "plus sign in a frame's out port tell you?",
     ["Threading links text frames so that a single story flows continuously from one frame into "
      "the next; the frames may be on the same page, the same spread or different pages.",
      "To thread: select a frame with the Selection tool, click its out port (the small square at "
      "the bottom right), then click inside the next frame — or drag to draw a new one.",
      "A red plus sign in the out port means OVERSET TEXT — copy that belongs to the story but has "
      "no frame to sit in.",
      "Overset text does not print and does not appear on screen, so it is the most common "
      "production error in a junior designer's file.",
      "Fix it by threading another frame, enlarging the existing frame, or editing the copy.",
      "Autoflow (Shift-click with the loaded cursor) adds both frames and pages until the whole "
      "story is placed.",
      "Taught in Topic 2 (Threading: One Story, Many Frames) and practised in Activity 6."]),

    ("K2",
     "Drawing curves is a core computer-aided drawing technique shared across InDesign and "
     "Illustrator.",
     "Which tool do you use to draw a curved path in InDesign, and how does the way you use it "
     "differ between a corner point and a smooth point?",
     ["The Pen tool (P) draws paths made of anchor points joined by segments.",
      "CLICK to place a corner point — this produces a straight segment to the previous point.",
      "CLICK AND DRAG to place a smooth point — dragging pulls out direction handles, and the angle "
      "and length of those handles determine the shape and size of the adjoining curve.",
      "Alt/Option-dragging a direction handle breaks its symmetry, letting a curve run straight "
      "into a corner.",
      "The Pencil tool draws freehand paths and the Smooth tool refines them; both are used for "
      "organic shapes rather than precision work.",
      "Existing points are edited with the Direct Selection tool (A), which selects the points "
      "themselves rather than the frame.",
      "Taught in Topic 2 (Drawing with the Pen Tool) and practised in Activity 8."]),

    ("K3",
     "Conventional signs and markings tell everyone downstream how the drawing is to be produced.",
     "Name the conventional signs and markings that appear on an InDesign page and on the PDF you "
     "send to print, and explain what each one signals.",
     ["On screen in InDesign: the black trim edge, the red bleed guide, the magenta margin guide, "
      "violet column guides, and cyan ruler guides — all non-printing.",
      "The baseline grid and document grid, which show the rhythm the text and objects align to.",
      "On the exported PDF: crop marks showing where the sheet is cut, registration marks aligning "
      "the four plates, colour bars for ink density, and the slug carrying the job name and date.",
      "Text-wrap boundaries and the frame's in/out ports signal how the story flows.",
      "They are produced by ticking Crop Marks and Use Document Bleed Settings under Marks and "
      "Bleeds when exporting Adobe PDF (Print).",
      "In manual drawing these marks were ruled by hand; in InDesign they are generated "
      "automatically and exactly, which removes measurement error but still requires the designer "
      "to set the correct values.",
      "Taught in Topic 1 (Anatomy of a Print Page) and Topic 3 (export), practised in Activity 19."]),
]

# ---------------------------------------------------------------- PRACTICAL PERFORMANCE (ABILITY)
SCENARIO = (
    "Harmony Petals is a Singapore florist with three retail outlets. Its marketing manager, "
    "Priya, has commissioned a two-page magazine spread to launch the new season's range, to be "
    "printed by Sun Ray Printers.\n\n"
    "The brief: A4 portrait pages, facing pages, full colour, litho-printed on 150 gsm gloss art "
    "paper. The hero photograph must run off the outer and top edges of the left-hand page. The "
    "body copy sits on the right-hand page in two columns, opening with a drop cap. Sun Ray "
    "Printers requires 3 mm bleed on all four sides, a 5 mm minimum safety margin from trim, and "
    "supply as PDF/X-1a with crop marks.\n\n"
    "The image and text assets are in the assessment folder supplied by your assessor. Work "
    "individually and take a screenshot at the end of each task, pasting it into the box provided."
)

# (label, codes, prompt, screenshot caption, model answer lines)
PRACTICAL = [
    ("Task 1", "A1, A2",
     "Establish the drawing and layout requirements for this spread, then create the InDesign "
     "document that meets them.\n\n"
     "Identify every element the layout must carry and every production specification the brief "
     "states — and name the ones the brief leaves out that you would have to confirm with the "
     "printer. Then determine the precise document dimensions, the finished size, the bleed and "
     "the margins, and create the document to those figures. Explain the process you followed in "
     "InDesign to arrive at the correct set-up.",
     "Take a screenshot at the end of each point and paste it in the box below: (A1, A2)",
     ["ESTABLISH THE REQUIREMENT (A1)",
      "Stated in the brief: A4 portrait, facing pages, full colour, litho, 150 gsm gloss art, hero "
      "photo bleeding off the outer and top edges, two-column body copy with a drop cap.",
      "Stated by the printer: 3 mm bleed all round, 5 mm minimum safety margin, PDF/X-1a with crop marks.",
      "To confirm before starting: the exact column gutter, the typeface and licensing, the spine/gutter "
      "allowance, and the deadline for supply. These are not in the brief and must not be assumed.",
      "Colour mode follows the output: litho press means CMYK, with images at 300 ppi effective at final size.",
      "",
      "DETERMINE THE DIMENSIONS (A2)",
      "A4 portrait trim = 210 x 297 mm. With 3 mm bleed on all four sides the artwork area is 216 x 303 mm.",
      "Facing pages ON, so the spread measures 420 x 297 mm at trim.",
      "",
      "CREATE THE DOCUMENT",
      "1. Close all documents, then Preferences > Units & Increments and set Horizontal and Vertical to Millimeters.",
      "2. File > New > Document, and choose the Print intent so the colour mode defaults to CMYK.",
      "3. Set Width 210 mm, Height 297 mm, Orientation Portrait, Pages 2, and TICK Facing Pages.",
      "4. Set the margins to at least 5 mm — commonly 15 mm top/outer/bottom with a wider inner "
      "   margin for the binding. Set Columns to 2 with a gutter of about 4-5 mm.",
      "5. Expand Bleed and Slug and set Bleed to 3 mm on all four sides. Click Create.",
      "6. Confirm the red bleed guide, black trim edge and magenta margin guides are all visible.",
      "7. File > Save As and name the file. Layout > Create Guides may be used to add a modular grid.",
      "",
      "EVIDENCE: a screenshot of the New Document dialog showing 210 x 297 mm, facing pages and 3 mm "
      "bleed, plus the created spread showing the bleed and margin guides.",
      "Corresponds to Activities 1, 2 and 4 in the Learner Guide."]),

    ("Task 2", "A3",
     "Identify and select the appropriate mediums, assets and conventional markings for this job.\n\n"
     "Examine the supplied assets and the document you created. Identify the conventional signs and "
     "markings present in the layout and explain what each one signals. Select and place the "
     "appropriate image and text mediums for a litho-printed magazine spread, justifying each "
     "choice of file format and resolution. Compare how these conventions are applied in a manual "
     "drawing against a digital InDesign layout, and discuss the advantages and the limitations of "
     "each method.",
     "Take a screenshot at the end of each point and paste it in the box below: (A3)",
     ["IDENTIFY THE CONVENTIONAL SIGNS AND MARKINGS",
      "On screen: black trim edge, red bleed guide, magenta margin guide, violet column guides, "
      "cyan ruler guides, the baseline grid, and text-frame in/out ports. None of these print.",
      "On output: crop marks, registration marks, colour bars and the slug — added at export under "
      "Marks and Bleeds.",
      "",
      "SELECT THE APPROPRIATE MEDIUMS",
      "Images: use File > Place (Ctrl/Cmd+D) so the asset is LINKED, not embedded. For a litho "
      "magazine choose PSD or TIFF for photographs (lossless, supports CMYK and transparency); JPEG "
      "is acceptable if supplied at quality and 300 ppi. PNG and GIF are screen formats and are not "
      "used for press work.",
      "Logos and any line artwork: place as AI, EPS or PDF so they stay vector and resolution-independent.",
      "Text: File > Place the supplied Word or RTF file rather than pasting, so structure is preserved "
      "and the source can be re-linked.",
      "",
      "VERIFY THE SELECTION IS CORRECT",
      "Open Window > Links and read Effective PPI for every placed image — it must be 300 ppi or "
      "better at the size used. Scaling a 300 ppi image to 200% halves it to 150 ppi, which is a fault.",
      "Confirm no link shows a red (missing) or yellow (modified) icon.",
      "Set Object > Fitting > Frame Fitting Options to Fill Frame Proportionally before placing.",
      "Extend the hero photograph past the trim so it fills the 3 mm bleed on the outer and top edges.",
      "",
      "MANUAL VERSUS DIGITAL",
      "Manual: marks are ruled by hand on a board; the advantage is a direct, tangible sense of the "
      "page, the limitations are measurement error, no undo, and re-drawing for every revision.",
      "Digital: marks are generated exactly and automatically, assets stay linked and updatable, and "
      "revisions cost minutes; the limitation is that InDesign will faithfully output whatever "
      "specification you gave it, so an incorrect setting is reproduced perfectly on 5,000 copies.",
      "",
      "EVIDENCE: a screenshot of the Links panel showing effective PPI and OK status, and of the "
      "spread with the guides visible and the photo bleeding off the trim.",
      "Corresponds to Activities 5, 6 and 7 in the Learner Guide."]),

    ("Task 3", "A4",
     "Refine the spread so that it meets the project requirements.\n\n"
     "Review the layout you have built. Identify the areas that fall short of the brief and make "
     "the adjustments needed — alignment and grid discipline, typography including the drop cap and "
     "the two-column body copy, text wrap around the hero image, and overall visual balance. "
     "Document each specific change you made and explain how it brings the spread closer to the "
     "brief. Finally, validate the file for press and produce the supply file.",
     "Take a screenshot at the end of each point and paste it in the box below: (A4)",
     ["REFINE THE TYPOGRAPHY",
      "Place the cursor in the opening paragraph and set the drop cap in the Paragraph panel — "
      "typically 3 lines, 1 character — and apply a character style to control its font and colour.",
      "Set the body copy in two columns via Object > Text Frame Options, with the gutter from the "
      "document set-up.",
      "Set leading to roughly 120% of the point size and align the body text to the baseline grid so "
      "the two columns line up across the spread.",
      "Adjust tracking and hyphenation to remove rivers, widows and orphans; use Optical kerning on "
      "the display type.",
      "",
      "REFINE THE LAYOUT",
      "Use the Align panel (Window > Object & Layout > Align) to align objects to the margins or the "
      "spread, then Distribute Spacing for even gaps.",
      "Apply Text Wrap (Window > Text Wrap) so the body copy flows around the hero image's shape "
      "rather than its rectangular frame; use Wrap Around Object Shape where the image is silhouetted.",
      "Check the stacking order with Object > Arrange, and separate background, images and text onto "
      "named layers in the Layers panel.",
      "Build paragraph, character and object styles for every repeating element so the change is "
      "reproducible rather than manual.",
      "",
      "VALIDATE AND OUTPUT",
      "Open Window > Output > Preflight and work to a profile checking missing links, image "
      "resolution below 300 ppi, overset text and RGB images in a CMYK job. Continue until the status "
      "bar reads 'No errors'.",
      "File > Package to collect the document, its links and its fonts into one folder for the printer.",
      "File > Export, Format Adobe PDF (Print), preset [PDF/X-1a:2001]. Under Marks and Bleeds tick "
      "Crop Marks and Use Document Bleed Settings.",
      "Open the exported PDF and confirm the artwork extends 3 mm past the crop marks.",
      "",
      "DOCUMENT THE CHANGES",
      "State each specific change and its justification — for example: 'Added a 3-line drop cap so "
      "the opening paragraph reads as the entry point'; 'Aligned body copy to the baseline grid so "
      "the two columns register across the spread'; 'Applied text wrap around the object shape so "
      "the copy follows the flower silhouette'; 'Replaced the 96 ppi image with the 300 ppi original'.",
      "",
      "EVIDENCE: screenshots of the refined spread, the Preflight panel showing no errors, and the "
      "PDF/X export dialog showing the marks and bleed settings.",
      "Corresponds to Activities 14, 15, 19 and 20 in the Learner Guide."]),
]

# ---------------------------------------------------------------- coverage check
def check_coverage():
    ks = {c for c, *_ in ORAL}
    as_ = set()
    for _, codes, *_ in PRACTICAL:
        as_ |= {c.strip() for c in codes.split(",")}
    need_k = {"K1", "K2", "K3"}
    need_a = {"A1", "A2", "A3", "A4"}
    print("\nCoverage map")
    print("  OQ questions :", len(ORAL), " codes:", ", ".join(sorted(ks)))
    for i, (c, *_rest) in enumerate(ORAL, 1):
        print(f"    Q{i} -> {c}")
    print("  PP tasks     :", len(PRACTICAL), " codes:", ", ".join(sorted(as_)))
    for i, (lbl, codes, *_rest) in enumerate(PRACTICAL, 1):
        print(f"    {lbl} -> {codes}")
    missing_k, missing_a = need_k - ks, need_a - as_
    if missing_k or missing_a:
        raise SystemExit(f"COVERAGE FAILURE — missing K: {sorted(missing_k)}  missing A: {sorted(missing_a)}")
    print("  OK — every K is covered by the OQ and every A by the PP.\n")

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, height_pt=90):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(1)
            if not ln.strip():
                b.add_run(" ").font.size = Pt(5); continue
            # A bare upper-case line is a sub-heading inside the model answer.
            if ln.isupper() or (ln.rstrip().endswith(")") and ln.split("(")[0].isupper()):
                rr = b.add_run(ln); rr.bold = True; rr.font.size = Pt(10.5); rr.font.color.rgb = BRAND
            else:
                b.paragraph_format.left_indent = Inches(0.15)
                rr = b.add_run("•  " + ln.strip()); rr.font.size = Pt(10.5)
    else:
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt * 20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]
LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text, oral=False):
    heading(doc, "Instructions to Candidate")
    items = ["This is an individual exercise.",
             "This is an open-book assessment.",
             f"A total of {minutes_text} is given to complete this assessment."]
    if oral:
        items += ["This assessment is conducted orally, one-to-one with the assessor. "
                  "Answer each question aloud in your own words.",
                  "The assessor may ask you to clarify or expand on an answer.",
                  "Your responses will be recorded or noted by the assessor.",
                  None]
    else:
        items += [None]
    items += BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            if oral:
                p.add_run(f"{i}.  Your course slides and Learner Guide are available for "
                          "reference during this open-book assessment at ").font.size = Pt(11)
            else:
                p.add_run(f"{i}.  Complete your answers on the document provided and "
                          "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

def title_block(doc, subtitle):
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, subtitle, size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

# ---------------------------------------------------------------- builders
def build_oq(answers):
    doc = base_doc()
    kind = "Oral Questioning (OQ) — Answer Key" if answers else "Oral Questioning (OQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    title_block(doc, "Answers to Oral Questioning (OQ)" if answers else "Oral Questioning (OQ)")
    if not answers:
        candidate_block(doc); instructions(doc, OQ_MINUTES, oral=True)
        grading(doc, "Candidate has answered all oral questions and demonstrated the underpinning "
                     "knowledge (K1, K2, K3) required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Oral Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer each question aloud, in your own words. Each question assesses underpinning "
              "knowledge taught in the course slides and the Learner Guide.",
         size=10.5, italic=True, color=GREY, after=8)
    per_page = 1 if answers else 2
    for i, (code, ctx, q, pts) in enumerate(ORAL, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({code})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None, height_pt=110)
        if i % per_page == 0 and i < len(ORAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answers to Oral Questioning (OQ) - {TITLE} - {suffix}.docx" if answers
            else f"Oral Questioning (OQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    title_block(doc, "Answers to Practical Performance Assessment" if answers
                     else "Practical Performance Assessment")
    if not answers:
        candidate_block(doc); instructions(doc, PP_MINUTES)
        grading(doc, "Candidate has successfully completed all the tasks for PP and is able to "
                     "explain the overall functions and features used to achieve these tasks.")
        page_break(doc)
    para(doc, "Practical Performance", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    for chunk in SCENARIO.split("\n\n"):
        para(doc, chunk, size=11, after=6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    for i, (label, codes, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({codes}):", size=11.5, bold=True, after=3, before=6)
        for chunk in prompt.split("\n\n"):
            para(doc, chunk, size=11, after=4)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, lines=pts if answers else None, height_pt=170)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set for", TITLE)
    check_coverage()
    build_pp(answers=False); build_pp(answers=True)
    build_oq(answers=False); build_oq(answers=True)
    print(f"Done. PP: {len(PRACTICAL)} tasks · OQ: {len(ORAL)} questions.")
