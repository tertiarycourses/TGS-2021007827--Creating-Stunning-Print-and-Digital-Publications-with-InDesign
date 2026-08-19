#!/usr/bin/env python3
"""Build the Learner Guide (LG) for
   Creating Stunning Print and Digital Publications with InDesign (TGS-2021007827)
as BOTH a Word document and an aligned Markdown mirror, from the single source.

The LG is where the DETAILED STEP-BY-STEP procedures live — the slide deck
deliberately carries none. Each of the 20 activities gets: the client scenario,
the objective, what you will produce, the numbered procedure with the exact menu
paths, the verification test, troubleshooting, and the discussion questions.

Also writes labs/activity-NN-*.md so each activity has its own folder-ready file.

Run:  python3 build_learner_guide.py
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import prodoc
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT = DOMAIN1 + DOMAIN2 + DOMAIN3

def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "labs")):
            return d
    return os.path.dirname(os.path.dirname(start))
REPO   = _find_repo(HERE)
ASSETS = os.path.join(os.path.dirname(HERE), "assets")
SCREENS = os.path.join(REPO, "courseware", "assets", "screens")
GEN     = os.path.join(REPO, "courseware", "assets", "gen")

HEADER_FILL = "1F6FEB"; TOPIC_FILL = "E8F0FE"; NOTE_FILL = "FFF7E8"
OK_FILL = "EEF7F3"; GREY = RGBColor(0x55, 0x5B, 0x66)
BRAND = RGBColor(0x1F, 0x6F, 0xEB)

def strip_md(s):
    return s.replace("**", "").replace("*", "")

def slug(s):
    s = re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")
    return re.sub(r"-+", "-", s)

def _img(name):
    for base in (GEN, SCREENS):
        for ext in (".png", ".jpg", ".jpeg"):
            p = os.path.join(base, name + ext)
            if os.path.exists(p): return p
    return None

# Concept figures embedded into the topic introductions of the LG.
TOPIC_FIGS = {
 1: [("page_anatomy", "Figure 1.1 — The anatomy of a print page: trim, bleed, margin and slug."),
     ("colour_modes", "Figure 1.2 — Colour mode follows the output: CMYK for ink, RGB for screen."),
     ("resolution_ladder", "Figure 1.3 — Resolution follows the medium; effective PPI is what prints."),
     ("workspace_map", "Figure 1.4 — The InDesign workspace."),
     ("frame_content", "Figure 1.5 — The frame and its content are two different things.")],
 2: [("threading", "Figure 2.1 — Threading: one story flowing through several frames."),
     ("app_roles", "Figure 2.2 — Which Creative Cloud application owns which job.")],
 3: [("styles_cascade", "Figure 3.1 — The styles cascade: change once, update everywhere."),
     ("preflight_gate", "Figure 3.2 — The preflight gate before every export."),
     ("export_matrix", "Figure 3.3 — Choosing the export format by its destination.")],
}

# Per-activity figure (the source artwork, imported from the original courseware)
ACT_IMG = {
 1:"act_masterpage", 2:"start_workspace", 3:"page_number_marker", 4:"act_grids",
 5:"act_textframe", 6:"act_threading", 7:"act_path", 8:"pen_curves",
 9:"pathfinder_panel", 10:"act_colour", 11:"transparency_effects", 12:"act_layers",
 13:"act_typepath", 14:"act_dropcap", 15:"nested_styles", 16:"act_table",
 17:"act_hyperlink", 18:"act_animation", 19:"preflight_gate", 20:"act_epub",
}

# ================================================================= DOCX
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc, "Learner Guide", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
 ("11.0", "1 September 2021",
  "Legacy release — learner guide slides covering 3 topics and 14 in-class activities.", C.TRAINER),
 (C.VERSION.lstrip("v"), C.VERSION_DATE,
  "Major revision. Rewritten as a standalone step-by-step Learner Guide against the current "
  f"InDesign release and the TSC {C.TSC_CODE} ability and knowledge statements. Expanded to 20 "
  "case-study activities on a single running client brief (Harmony Petals), each with a full "
  "numbered procedure, verification test, troubleshooting notes and discussion questions. New "
  "material on establishing the drawing requirement, effective resolution and the Links panel, "
  "compound paths and Pathfinder, object styles, preflight and package, PDF/X export, and "
  "reflowable versus fixed-layout EPUB.", C.TRAINER),
])
prodoc.add_toc(doc)

def H(t, level=1): return doc.add_heading(t, level=level)

def para(t, size=11, bold=False, italic=False, color=None, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(t, size=10.5):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(t).font.size = Pt(size); return p

def set_cell(cell, text, bold=False, size=9.5, color=None, fill=None):
    cell.text = ""; p = cell.paragraphs[0]
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
    if color: r.font.color.rgb = color
    if fill: prodoc._shade_cell(cell, fill)

def callout(label, text, fill=NOTE_FILL):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(label.upper() + "  "); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = BRAND
    r2 = p.add_run(text); r2.font.size = Pt(10)
    prodoc._shade_cell(c, fill)
    doc.add_paragraph()
    return t

def figure(name, caption, width=6.0):
    p = _img(name)
    if not p: return
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run().add_picture(p, width=Inches(width))
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

# ---------------------------------------------------------------- front matter
H("About This Course", 1)
para(f"{C.TITLE} is a {C.DAYS}-day ({C.HOURS_PER_DAY}-hour) WSQ course, reference "
     f"{C.COURSE_CODE}, delivered by {C.ORG}. It teaches you to take a publication from a "
     "client brief through to a file a commercial printer or a digital platform will accept "
     "without a query.")
para("This Learner Guide is your step-by-step reference. The slide deck used in class shows "
     "what each feature is for and why it matters; this guide gives you the exact procedure, "
     "menu path by menu path, so you can repeat every activity at your own desk afterwards. "
     "You may use this guide during the open-book assessment.")

H("How This Guide Is Organised", 2)
for b in ["Each topic opens with the concepts you need before you touch the software.",
          "Each activity begins with a real client situation, so you practise judgement, not just clicks.",
          "Every procedure is numbered, with the exact InDesign menu path in bold.",
          "Every activity ends with a verification test, troubleshooting notes and discussion questions.",
          "The discussion questions are the same kind of question asked in the Written Assessment."]:
    bullet(b)

H("The Running Case Study — Harmony Petals", 2)
para("Every activity in this guide is set at Harmony Petals, a Singapore florist with three "
     "retail outlets. You are its junior layout designer. Priya, the marketing manager, briefs "
     "you; Sun Ray Printers produces the printed work. The jobs build on one another exactly as "
     "they would in a real studio: a promotional flyer, a seasonal care guide, a quarterly "
     "magazine and a digital edition. The stakes named in each brief — reprint costs, missed "
     "press slots, rejected files — are the ones that actually arise in Singapore print production.")

H("Learning Outcomes", 1)
para("On completion of this course, you will be able to:")
for t, b in C.LEARNING_OUTCOMES_LONG:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t + " — "); r.bold = True; r.font.size = Pt(10.5)
    p.add_run(b).font.size = Pt(10.5)

H("Skills Framework Alignment", 1)
para(f"TSC Title: {C.TSC_TITLE}    ·    TSC Code: {C.TSC_CODE}")
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
hdr = t.add_row().cells
for i, h in enumerate(["Ref", "Statement"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for ref, stmt in C.TSC_ABILITIES + C.TSC_KNOWLEDGE:
    c = t.add_row().cells
    set_cell(c[0], ref, bold=True, size=9.5, fill=TOPIC_FILL); set_cell(c[1], stmt, size=9.5)
    c[0].width = Inches(0.7); c[1].width = Inches(6.1)

H("Before You Start", 1)
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
for k, v in C.PREREQUISITES:
    c = t.add_row().cells
    set_cell(c[0], k, bold=True, size=10, fill=TOPIC_FILL); set_cell(c[1], v, size=10)
    c[0].width = Inches(1.6); c[1].width = Inches(5.2)
doc.add_paragraph()
callout("Tip", "Download the lab files from https://lms-tms.tertiaryinfotech.com and unzip them to "
               "your desktop before the class begins. Activity 1 needs them.")

H("Assessment", 1)
for a in [C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
          "Format: Open Book — you may use the course slides, this Learner Guide and approved materials only.",
          "You are graded Competent (C) or Not Yet Competent (NYC) on every ability and knowledge statement. A single NYC means NYC for the whole unit.",
          C.ASSESSMENT["note"]]:
    bullet(a)

# ---------------------------------------------------------------- topics
for tp in C.TOPICS:
    doc.add_page_break()
    H(f"Topic {tp['code']} — {tp['title']}", 1)
    para(tp["subtitle"], italic=True, color=GREY)
    para(f"Topic weighting: {tp['weighting']} of the course.", size=10, color=GREY)

    H("Key Concepts", 2)
    for con in tp["concepts"]:
        bullet(con, size=10.5)

    for fig, cap in TOPIC_FIGS.get(tp["num"], []):
        figure(fig, cap)

    acts = [a for a in ACT if a["topic"] == tp["num"]]
    for a in acts:
        doc.add_page_break()
        H(f"Activity {a['num']} — {a['title']}", 2)

        # scenario
        H("The Situation", 3)
        para(strip_md(a["scenario"]), size=10.5)

        # objective / build / tools
        t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
        for k, v in [("Learning outcome", a["lo"]),
                     ("Objective", a["objective"]),
                     ("You will produce", a["build"]),
                     ("Tools and panels", a["services"])]:
            c = t.add_row().cells
            set_cell(c[0], k, bold=True, size=9.5, fill=TOPIC_FILL); set_cell(c[1], v, size=9.5)
            c[0].width = Inches(1.5); c[1].width = Inches(5.3)
        doc.add_paragraph()

        img = ACT_IMG.get(a["num"])
        if img:
            figure(img, f"Figure {a['num']} — Reference artwork for Activity {a['num']}.", width=4.6)

        H("What You Will Do", 3)
        para(a["desc"], size=10.5)

        # THE STEP-BY-STEP PROCEDURE — the heart of the LG
        H("Step-by-Step Procedure", 3)
        for i, (text, cmd) in enumerate(a["steps"], 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            r = p.add_run(f"{i}.  "); r.bold = True; r.font.size = Pt(10.5)
            p.add_run(text).font.size = Pt(10.5)
            if cmd:
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Inches(0.6)
                cp.paragraph_format.space_after = Pt(7)
                cr = cp.add_run(cmd); cr.bold = True; cr.font.size = Pt(9.5); cr.font.color.rgb = BRAND

        H("Verify Your Work", 3)
        callout("Done when", a["test"], fill=OK_FILL)

        H("If It Doesn't Work", 3)
        para(a["troubleshoot"], size=10.5)

        H("Discussion Questions", 3)
        para("Answer these in your own words and aloud. The Oral Questioning assessment asks the same kind of question.",
             size=10, italic=True, color=GREY)
        for qi, q in enumerate(a["questions"], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            r = p.add_run(f"Q{qi}.  "); r.bold = True; r.font.size = Pt(10.5)
            p.add_run(q).font.size = Pt(10.5)
        doc.add_paragraph()

    H(f"Topic {tp['code']} Recap", 2)
    for a in acts:
        bullet("You can now " + a["objective"].split("(TSC")[0].strip().rstrip(".").lower() + ".", size=10.5)

# ---------------------------------------------------------------- back matter
doc.add_page_break()
H("Keyboard Shortcuts Worth Knowing", 1)
t = doc.add_table(rows=0, cols=3); t.style = "Table Grid"
hdr = t.add_row().cells
for i, h in enumerate(["Action", "Windows", "macOS"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
SHORTCUTS = [
 ("Selection tool (frame)", "V", "V"), ("Direct Selection tool (content / points)", "A", "A"),
 ("Type tool", "T", "T"), ("Pen tool", "P", "P"), ("Rectangle Frame tool", "F", "F"),
 ("Preview / Normal view toggle", "W", "W"), ("Place a file", "Ctrl+D", "Cmd+D"),
 ("Save", "Ctrl+S", "Cmd+S"), ("Undo", "Ctrl+Z", "Cmd+Z"),
 ("Fit page in window", "Ctrl+0", "Cmd+0"), ("Actual size (100%)", "Ctrl+1", "Cmd+1"),
 ("Show / hide guides", "Ctrl+;", "Cmd+;"), ("Lock guides", "Ctrl+Alt+;", "Cmd+Opt+;"),
 ("Show / hide baseline grid", "Ctrl+Alt+'", "Cmd+Opt+'"),
 ("Group / Ungroup", "Ctrl+G / Ctrl+Shift+G", "Cmd+G / Cmd+Shift+G"),
 ("Paste in place", "Ctrl+Alt+Shift+V", "Cmd+Opt+Shift+V"),
 ("Export", "Ctrl+E", "Cmd+E"), ("Package", "Ctrl+Alt+Shift+P", "Cmd+Opt+Shift+P"),
 ("Character panel", "Ctrl+T", "Cmd+T"), ("Paragraph panel", "Ctrl+Alt+T", "Cmd+Opt+T"),
]
for a, w, m in SHORTCUTS:
    c = t.add_row().cells
    set_cell(c[0], a, size=9.5); set_cell(c[1], w, size=9.5); set_cell(c[2], m, size=9.5)
    c[0].width = Inches(3.4); c[1].width = Inches(1.7); c[2].width = Inches(1.7)

doc.add_paragraph()
H("Glossary", 1)
GLOSSARY = [
 ("Bleed", "Artwork extended beyond the trim edge (3 mm in Singapore/ISO practice) so that a small cutting drift never leaves a white sliver."),
 ("Trim", "The finished, cut size of the page. This is the 'page size' entered in the New Document dialog."),
 ("Slug", "The area outside the bleed carrying job name, colour bars and printer's marks; discarded after trimming."),
 ("Live matter", "Text, logos and any content that must survive the cut — kept inside the margin."),
 ("Effective PPI", "The resolution of a placed image after it has been scaled in InDesign. This, not the file's native PPI, is what prints."),
 ("Overset text", "Copy that belongs to a story but has no frame to sit in. Marked by a red plus in the out port; it never prints."),
 ("Threading", "Linking text frames so one story flows between them."),
 ("Parent page", "A reusable page background (formerly 'master page') carrying repeating elements applied across many pages."),
 ("Compound path", "Two or more paths combined so that overlapping areas become transparent holes."),
 ("Clipping path", "A vector outline that crops an image so only part of it shows."),
 ("Spot colour", "A pre-mixed ink (e.g. Pantone) printed from its own plate, used for brand-critical colour."),
 ("Process colour", "Colour built from the four CMYK process inks."),
 ("Preflight", "The continuous check of a live document against a profile — links, resolution, overset text, fonts, colour."),
 ("Package", "File > Package — collects the document, every linked asset and every used font into one folder for hand-off."),
 ("PDF/X", "An ISO subset of PDF for graphic-arts exchange. X-1a is CMYK and flattened; X-4 supports live transparency and colour management."),
 ("Reflowable EPUB", "An e-book whose text re-flows to fit the reader's device and chosen type size."),
 ("Fixed-layout EPUB", "An e-book that preserves the designed layout exactly, page for page."),
 ("Style (paragraph / character / object)", "A named, reusable set of formatting attributes. Redefine the style and every element using it updates."),
 ("Nested style", "A character style applied automatically within a paragraph up to a defined character."),
 ("Baseline grid", "A non-printing horizontal grid that locks body text to a common rhythm across columns and pages."),
]
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
hdr = t.add_row().cells
for i, h in enumerate(["Term", "Meaning"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for term, mean in GLOSSARY:
    c = t.add_row().cells
    set_cell(c[0], term, bold=True, size=9.5, fill=TOPIC_FILL); set_cell(c[1], mean, size=9.5)
    c[0].width = Inches(1.9); c[1].width = Inches(4.9)

doc.add_paragraph()
H("Where to Go Next", 1)
for r in ["Adobe InDesign Learn & Support — https://helpx.adobe.com/indesign/desktop.html",
          "Get started with InDesign — https://www.adobe.com/learn/indesign/web/get-started-indesign",
          "InDesignSkills tutorials and templates — https://indesignskills.com",
          "Adobe InDesign product page — https://www.adobe.com/sg/products/indesign.html",
          "Free trial — https://www.adobe.com/sg/products/indesign/free-trial-download.html",
          "InDesign for developers (scripting and UXP) — https://developer.adobe.com/indesign/"]:
    bullet(r)

doc.add_paragraph()
H("Support", 1)
for r in ["Email: enquiry@tertiaryinfotech.com", "Tel: +65 6100 0613",
          "Website: www.tertiarycourses.com.sg",
          "LMS / TMS: https://lms-tms.tertiaryinfotech.com"]:
    bullet(r)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT = os.path.join(REPO, "courseware", f"LG-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved", OUT)

# ================================================================= MARKDOWN MIRROR
md = []
md.append(f"# {C.TITLE}\n")
md.append(f"**Learner Guide** · Version {C.VERSION} · {C.VERSION_DATE}\n")
md.append(f"WSQ Course Reference: **{C.COURSE_CODE}** · {C.ORG} · {C.UEN}\n")
md.append(f"\n> Skills Framework: {C.TSC_TITLE} ({C.TSC_CODE})\n")

md.append("\n## About This Course\n")
md.append(f"{C.TITLE} is a {C.DAYS}-day ({C.HOURS_PER_DAY}-hour) WSQ course, reference "
          f"{C.COURSE_CODE}, delivered by {C.ORG}. It teaches you to take a publication from a "
          "client brief through to a file a commercial printer or a digital platform will accept "
          "without a query.\n")
md.append("\nThis Learner Guide is your step-by-step reference. The slide deck used in class shows "
          "what each feature is for and why it matters; this guide gives you the exact procedure, "
          "menu path by menu path. You may use this guide during the open-book assessment.\n")

md.append("\n### The Running Case Study — Harmony Petals\n")
md.append("Every activity is set at **Harmony Petals**, a Singapore florist with three retail "
          "outlets. You are its junior layout designer. **Priya**, the marketing manager, briefs "
          "you; **Sun Ray Printers** produces the printed work.\n")

md.append("\n## Learning Outcomes\n")
for t_, b in C.LEARNING_OUTCOMES_LONG:
    md.append(f"- **{t_}** — {b}\n")

md.append("\n## Skills Framework Alignment\n\n| Ref | Statement |\n|---|---|\n")
for ref, stmt in C.TSC_ABILITIES + C.TSC_KNOWLEDGE:
    md.append(f"| **{ref}** | {stmt} |\n")

md.append("\n## Before You Start\n\n| | |\n|---|---|\n")
for k, v in C.PREREQUISITES:
    md.append(f"| **{k}** | {v} |\n")

md.append("\n## Assessment\n")
for a in [C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, this Learner Guide and approved materials only.",
          "Graded Competent (C) or Not Yet Competent (NYC) on every ability and knowledge statement.",
          C.ASSESSMENT["note"]]:
    md.append(f"- {a}\n")

for tp in C.TOPICS:
    md.append(f"\n---\n\n## Topic {tp['code']} — {tp['title']}\n")
    md.append(f"*{tp['subtitle']}*\n\nTopic weighting: {tp['weighting']}\n")
    md.append("\n### Key Concepts\n")
    for con in tp["concepts"]:
        md.append(f"- {con}\n")
    for fig, cap in TOPIC_FIGS.get(tp["num"], []):
        md.append(f"\n![{cap}](courseware/assets/gen/{fig}.png)\n\n*{cap}*\n")

    for a in [x for x in ACT if x["topic"] == tp["num"]]:
        md.append(f"\n### Activity {a['num']} — {a['title']}\n")
        md.append(f"\n**The situation.** {a['scenario']}\n")
        md.append(f"\n| | |\n|---|---|\n")
        md.append(f"| **Learning outcome** | {a['lo']} |\n")
        md.append(f"| **Objective** | {a['objective']} |\n")
        md.append(f"| **You will produce** | {a['build']} |\n")
        md.append(f"| **Tools and panels** | {a['services']} |\n")
        md.append(f"\n**What you will do.** {a['desc']}\n")
        md.append("\n#### Step-by-Step Procedure\n\n")
        for i, (text, cmd) in enumerate(a["steps"], 1):
            md.append(f"{i}. {text}\n")
            if cmd:
                md.append(f"   > `{cmd}`\n")
        md.append(f"\n#### Verify Your Work\n\n> ✅ **Done when:** {a['test']}\n")
        md.append(f"\n#### If It Doesn't Work\n\n{a['troubleshoot']}\n")
        md.append("\n#### Discussion Questions\n\n")
        for qi, q in enumerate(a["questions"], 1):
            md.append(f"{qi}. {q}\n")

    md.append(f"\n### Topic {tp['code']} Recap\n")
    for a in [x for x in ACT if x["topic"] == tp["num"]]:
        md.append("- You can now " + a["objective"].split("(TSC")[0].strip().rstrip(".").lower() + ".\n")

md.append("\n---\n\n## Keyboard Shortcuts Worth Knowing\n\n| Action | Windows | macOS |\n|---|---|---|\n")
for a, w, m in SHORTCUTS:
    md.append(f"| {a} | `{w}` | `{m}` |\n")

md.append("\n## Glossary\n\n| Term | Meaning |\n|---|---|\n")
for term, mean in GLOSSARY:
    md.append(f"| **{term}** | {mean} |\n")

md.append("\n## Where to Go Next\n")
for r in ["[Adobe InDesign Learn & Support](https://helpx.adobe.com/indesign/desktop.html)",
          "[Get started with InDesign](https://www.adobe.com/learn/indesign/web/get-started-indesign)",
          "[InDesignSkills](https://indesignskills.com)",
          "[Adobe InDesign product page](https://www.adobe.com/sg/products/indesign.html)",
          "[Free trial](https://www.adobe.com/sg/products/indesign/free-trial-download.html)",
          "[InDesign for developers](https://developer.adobe.com/indesign/)"]:
    md.append(f"- {r}\n")

md.append("\n## Support\n")
md.append("- Email: enquiry@tertiaryinfotech.com\n- Tel: +65 6100 0613\n"
          "- Website: www.tertiarycourses.com.sg\n- LMS / TMS: https://lms-tms.tertiaryinfotech.com\n")
md.append(f"\n---\n\n© 2026 {C.ORG}. All rights reserved.\n")

MDOUT = os.path.join(REPO, f"LG-{C.SHORT_TITLE}.md")
open(MDOUT, "w").write("".join(md))
print("Saved", MDOUT)

# ================================================================= PER-ACTIVITY FILES
# Each activity gets its OWN folder under activities/, with its own markdown file,
# as required by the course brief.
ACTROOT = os.path.join(REPO, "activities")
os.makedirs(ACTROOT, exist_ok=True)
index = ["# Activities — " + C.TITLE + "\n",
         f"\nWSQ Course Reference: **{C.COURSE_CODE}** · Version {C.VERSION} · {C.VERSION_DATE}\n",
         "\nEvery activity has its own folder containing a full step-by-step brief.\n",
         "\n| # | Activity | Topic | LO | Folder |\n|---|---|---|---|---|\n"]
for a in ACT:
    name = f"activity-{a['num']:02d}-{slug(a['title'])}"
    folder = os.path.join(ACTROOT, name)
    os.makedirs(folder, exist_ok=True)
    tp = next(t for t in C.TOPICS if t["num"] == a["topic"])
    out = [f"# Activity {a['num']} — {a['title']}\n",
           f"\n**Topic {tp['code']}: {tp['title']}**  ·  **{a['lo']}**  ·  {C.COURSE_CODE}\n",
           f"\n## The Situation\n\n{a['scenario']}\n",
           "\n## At a Glance\n\n| | |\n|---|---|\n",
           f"| **Learning outcome** | {a['lo']} |\n",
           f"| **Objective** | {a['objective']} |\n",
           f"| **You will produce** | {a['build']} |\n",
           f"| **Tools and panels** | {a['services']} |\n",
           f"\n## What You Will Do\n\n{a['desc']}\n",
           "\n## Step-by-Step Procedure\n\n"]
    for i, (text, cmd) in enumerate(a["steps"], 1):
        out.append(f"{i}. {text}\n")
        if cmd:
            out.append(f"   > `{cmd}`\n")
    out.append(f"\n## Verify Your Work\n\n> ✅ **Done when:** {a['test']}\n")
    out.append(f"\n## If It Doesn't Work\n\n{a['troubleshoot']}\n")
    out.append("\n## Discussion Questions\n\n")
    for qi, q in enumerate(a["questions"], 1):
        out.append(f"{qi}. {q}\n")
    img = ACT_IMG.get(a["num"])
    if img:
        out.append(f"\n## Reference Artwork\n\n![Activity {a['num']}](../../courseware/assets/screens/{img}.png)\n")
    out.append(f"\n---\n\n© 2026 {C.ORG}. All rights reserved.\n")
    open(os.path.join(folder, "README.md"), "w").write("".join(out))
    index.append(f"| {a['num']} | {a['title']} | {tp['code']} | {a['lo']} | [`{name}/`]({name}/) |\n")

open(os.path.join(ACTROOT, "README.md"), "w").write("".join(index))
print("Saved", len(ACT), "activity folders under", ACTROOT)
