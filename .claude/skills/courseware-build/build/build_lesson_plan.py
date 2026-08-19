#!/usr/bin/env python3
"""Build the Lesson Plan (LP) for
   Creating Stunning Print and Digital Publications with InDesign (TGS-2021007827).

One training day of EXACTLY 8 instructional hours (480 min), 9:30 am – 6:30 pm with a
1-hour lunch; tea breaks are counted inside training time. Every schedule row names the
activity AND the slide range in the current deck, so the LP and the PPT stay aligned.

Run:  python3 build_lesson_plan.py
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import prodoc
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT = DOMAIN1 + DOMAIN2 + DOMAIN3

def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "labs")):
            return d
    return os.path.dirname(os.path.dirname(start))
REPO   = _find_repo(HERE)
ASSETS = os.path.join(os.path.dirname(HERE), "assets")

HEADER_FILL = "1F6FEB"; TOPIC_FILL = "E8F0FE"; BREAK_FILL = "FFF7E8"
LUNCH_FILL  = "FDECEF"; ASSESS_FILL = "EEF7F3"
GREY = RGBColor(0x55, 0x5B, 0x66)

def acts(nums):
    return "; ".join(f"Activity {a['num']} {a['title']}" for a in ACT if a["num"] in nums)

# ---------------------------------------------------------------- slide references
# Slide numbers are DERIVED from the built deck (slide_index.json, written by
# build_slides.py) — never hand-typed, so the LP can never drift from the PPT.
import json
_IDX_PATH = os.path.join(HERE, "slide_index.json")
try:
    IDX = json.load(open(_IDX_PATH))
except Exception:
    raise SystemExit(f"{_IDX_PATH} is missing — run build_slides.py first so the LP can "
                     "derive its slide references from the actual deck.")

def sl(a, b=None):
    """Render a slide reference, e.g. '[Slides 26-40]' or '[Slide 25]'."""
    if b is None or b == a:
        return f"  [Slide {a}]"
    return f"  [Slides {a}-{b}]"

def act_span(nums, from_divider=None):
    """The slide range covering a set of activities (4 slides each in the deck).
    from_divider=<topic no> starts the range at that topic's 'Hands-On Activities'
    divider slide, so the divider is never orphaned outside every schedule row."""
    starts = [IDX["activity_slide"][str(n)] for n in sorted(nums)]
    lo = starts[0]
    if from_divider is not None:
        d = IDX.get(f"t{from_divider}_labs_start")
        if d and d < lo:
            lo = d
    return lo, starts[-1] + 3

# ---------------------------------------------------------------- the schedule
# (start, end, minutes, kind, text). Lunch is EXCLUDED from instructional hours;
# tea breaks are counted within them. Total instructional time must be 480 min.
DAY1 = [
 ("9:30","9:45",  15,"admin", "Digital Attendance (AM) · Welcome · Trainer and learner introductions · Ground rules" + sl(2, 7)),
 ("9:45","10:00", 15,"admin", "Course material download from the LMS · Skills Framework alignment (TSC " + C.TSC_CODE + ") · Learning outcomes · Course outline · Assessment briefing" + sl(8, IDX["core_concepts"] - 1)),
 ("10:00","10:30",30,"topic", "Core concepts: why InDesign · what it is for · which Creative Cloud application owns which job · InDesign vs the alternatives" + sl(IDX["core_concepts"], IDX["core_end"])),
 ("10:30","11:00",30,"topic", "Topic 1 Get Started on InDesign — key concepts: establishing the drawing and layout requirement, page anatomy (trim, bleed, margin, slug), colour mode and resolution" + sl(IDX["t1_section"], IDX["t1_concepts_start"] + 6)),
 ("11:00","11:15",15,"break", "Tea break (counted within training time)"),
 ("11:15","11:45",30,"topic", "Topic 1 continued — the InDesign interface and workspace, the New Document dialog, the two selection tools, parent pages, automatic page numbering, grids and guides, Adjust Layout" + sl(IDX["t1_concepts_start"] + 7, IDX["t1_concepts_end"])),
 ("11:45","13:00",75,"lab",   "Topic 1 hands-on — " + acts({1,2,3,4}) + sl(*act_span({1,2,3,4}, from_divider=1))),
 ("13:00","14:00",60,"lunch", "Lunch Break"),
 ("14:00","14:10",10,"admin", "Digital Attendance (PM) · Recap of the morning" + sl(IDX["t1_recap"])),
 ("14:10","14:45",35,"topic", "Topic 2 Basic InDesign Drawing Techniques — key concepts: frames and content, text and threading, placing graphics and the Links panel, paths and the Pen tool" + sl(IDX["t2_section"], IDX["t2_concepts_start"] + 7)),
 ("14:45","15:45",60,"lab",   "Topic 2 hands-on part 1 — " + acts({5,6,7,8}) + sl(*act_span({5,6,7,8}, from_divider=2))),
 ("15:45","16:00",15,"break", "Tea break (counted within training time)"),
 ("16:00","16:20",20,"topic", "Topic 2 continued — compound paths and Pathfinder, clipping paths and text wrap, colour and swatches, transparency and effects, transform, align and layers" + sl(IDX["t2_concepts_start"] + 8, IDX["t2_concepts_end"])),
 ("16:20","17:00",40,"lab",   "Topic 2 hands-on part 2 — " + acts({9,10,11,12}) + sl(*act_span({9,10,11,12}))),
 ("17:00","17:25",25,"topic", "Topic 3 Refine InDesign Drawings — key concepts: typography and type on a path, the styles cascade, tables, interactivity, preflight, package and export" + sl(IDX["t3_section"], IDX["t3_concepts_end"])),
 ("17:25","18:10",45,"lab",   "Topic 3 hands-on — " + acts({13,14,15,16,17,18,19,20}) + sl(*act_span({13,14,15,16,17,18,19,20}, from_divider=3))),
 ("18:10","18:20",10,"recap", "Course summary · what you achieved · continuing your learning · Q&A" + sl(IDX["t3_recap"], IDX["total"])),
 ("18:20","18:30",10,"assess","Course Feedback and TRAQOM Survey · Digital Attendance (Assessment)"),
]

# The assessment sits outside the 8 instructional hours, as stated on the course page.
ASSESS = [
 ("18:30","19:45",75,"assess","Practical Performance (PP) — hands-on InDesign layout tasks, 75 minutes, open book. Assesses A1, A2, A3, A4."),
 ("19:45","20:00",15,"assess","Oral Questioning (OQ) — five questions, one-to-one with the assessor, 15 minutes, open book. Assesses K1, K2, K3."),
]

SCHEDULE = {1: (C.DAY_THEMES[1], DAY1)}

# ---------------------------------------------------------------- document
doc = Document()
for _sec in doc.sections:                     # a little more width for the schedule table
    _sec.left_margin = Inches(0.85); _sec.right_margin = Inches(0.85)
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc, "LESSON PLAN", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
 ("11.0", "1 September 2021",
  "Legacy release — 3-topic lesson plan with 14 in-class activities.", C.TRAINER),
 (C.VERSION.lstrip("v"), C.VERSION_DATE,
  "Major revision. Rebuilt against the current InDesign release and the TSC "
  f"{C.TSC_CODE} ability and knowledge statements. Expanded from 14 to 20 case-study "
  "activities built on a single running client brief (Harmony Petals). Added establishing "
  "the drawing requirement, effective resolution, the Links panel, compound paths and "
  "Pathfinder, object styles, preflight, package and PDF/X export, and reflowable vs "
  "fixed-layout EPUB. Schedule re-timed to 9:30 am – 6:30 pm with a 1-hour lunch. "
  "Slide references updated to the v12.0 deck.", C.TRAINER),
])
prodoc.add_toc(doc)

def H(text, level=1):
    return doc.add_heading(text, level=level)

def set_cell(cell, text, bold=False, size=9.5, color=None, fill=None, align=None):
    cell.text = ""; p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
    if color: r.font.color.rgb = color
    if fill: prodoc._shade_cell(cell, fill)

def fix_widths(tbl, widths_in):
    """Pin a table to explicit column widths: fixed layout + a real tblGrid + a
    width on every cell. Without the grid, LibreOffice re-distributes the columns
    and squeezes the widest one."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _El
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(_qn('w:tblLayout')):
        tblPr.remove(old)
    lay = _El('w:tblLayout'); lay.set(_qn('w:type'), 'fixed'); tblPr.append(lay)
    for old in tbl._tbl.findall(_qn('w:tblGrid')):
        tbl._tbl.remove(old)
    grid = _El('w:tblGrid')
    for w in widths_in:
        gc = _El('w:gridCol'); gc.set(_qn('w:w'), str(int(w * 1440))); grid.append(gc)
    tbl._tbl.insert(1, grid)
    for row in tbl.rows:
        for c, w in zip(row.cells, widths_in):
            c.width = Inches(w)
    return tbl

def kv_table(rows):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in rows:
        c = t.add_row().cells
        set_cell(c[0], k, bold=True, size=10, fill=TOPIC_FILL)
        set_cell(c[1], v, size=10)
        c[0].width = Inches(2.1); c[1].width = Inches(4.7)
    return t

# ---- Course Information
H("Course Information", 1)
kv_table([
 ("Course Title", C.TITLE),
 ("WSQ Course Reference", C.COURSE_CODE),
 ("Training Provider", C.ORG + "  (" + C.UEN.replace("UEN: ", "UEN ") + ")"),
 ("Duration", f"{C.DAYS} day · {C.HOURS_PER_DAY} instructional hours, plus 90 minutes of assessment"),
 ("Daily Timing", "9:30 am – 6:30 pm (1-hour lunch break; tea breaks counted within training time)"),
 ("Mode of Delivery", "Instructor-led classroom delivery with hands-on Adobe InDesign practice at every learner's workstation"),
 ("Trainer", C.TRAINER),
 ("Skills Framework", f"{C.TSC_TITLE}  ({C.TSC_CODE})"),
])

# ---- Target audience & prerequisites
H("Target Audience", 1)
for a in C.TARGET_AUDIENCE:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10.5)

H("Prerequisites", 1)
kv_table(C.PREREQUISITES)

# ---- Learning outcomes
H("Learning Outcomes", 1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size = Pt(10.5)

# ---- Skills Framework mapping
H("Skills Framework Alignment", 1)
doc.add_paragraph(f"TSC Title: {C.TSC_TITLE}    ·    TSC Code: {C.TSC_CODE}")
t = doc.add_table(rows=0, cols=3); t.style = "Table Grid"
hdr = t.add_row().cells
for i, h in enumerate(["Ref", "TSC Ability / Knowledge Statement", "Covered by"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
COVER = {
 "A1": "Topic 1 · Activities 1, 3; Topic 3 · Activity 19",
 "A2": "Topic 1 · Activities 1–4",
 "A3": "Topic 1 · Activity 2; Topic 2 · Activities 7, 10; Topic 3 · Activities 19, 20",
 "A4": "Topic 2 · Activities 8–12; Topic 3 · Activities 13–20",
 "K1": "Topic 1 — drawing conventions and manual layout principles",
 "K2": "Topic 1 — InDesign as computer-aided layout software and its place in Creative Cloud",
 "K3": "Topic 2 and Topic 3 — the digital production drawing workflow end to end",
 "K4": "Topic 1 — trim, bleed, margins, grids, scales and finished sizes",
 "K5": "Topic 1 and Topic 3 — colour mode, resolution and output medium selection",
}
for ref, stmt in C.TSC_ABILITIES + C.TSC_KNOWLEDGE:
    c = t.add_row().cells
    set_cell(c[0], ref, bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(c[1], stmt, size=9.5)
    set_cell(c[2], COVER.get(ref, ""), size=9.5)
    c[0].width = Inches(0.6); c[1].width = Inches(3.6); c[2].width = Inches(2.6)

# ---- Assessment
H("Assessment", 1)
for a in [C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "The assessment is conducted at the end of the training day, after the TRAQOM survey "
          "and the Assessment digital attendance have been taken.",
          "Learners are graded Competent (C) or Not Yet Competent (NYC) on each instrument.",
          C.ASSESSMENT["note"]]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10.5)

# ---- Schedule
KIND_FILL = {"topic": TOPIC_FILL, "break": BREAK_FILL, "lunch": LUNCH_FILL,
             "assess": ASSESS_FILL, "admin": "F3F5F8", "recap": "F3F5F8", "lab": None}

H("Course Schedule", 1)
for day, (theme, rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}", 2)
    tbl = doc.add_table(rows=0, cols=3); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.add_row().cells
    for i, h in enumerate(["Time", "Duration", "Topic / Activity  [slide reference]"]):
        set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
    training = 0
    for start, end, mins, kind, text in rows:
        cells = tbl.add_row().cells; fill = KIND_FILL.get(kind)
        set_cell(cells[0], f"{start}–{end}", bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        set_cell(cells[1], f"{mins} min", size=9.5, fill=fill)
        set_cell(cells[2], text, bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        if kind != "lunch": training += mins
    fix_widths(tbl, [0.95, 0.7, 5.15])
    p = doc.add_paragraph()
    r = p.add_run("Note: Topic 2's key concepts are taught in two blocks either side of the "
                  "first practice session, so its slide references intentionally return to an "
                  "earlier range after the 14:45 activities.")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    p = doc.add_paragraph()
    r = p.add_run(f"Total instructional time: {training} minutes ({training // 60} hours), "
                  f"excluding the 60-minute lunch break.")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
    assert training == 480, f"Day {day} instructional minutes = {training}, expected 480"

H("Assessment Session (following the training day)", 2)
tbl = doc.add_table(rows=0, cols=3); tbl.style = "Table Grid"
hdr = tbl.add_row().cells
for i, h in enumerate(["Time", "Duration", "Assessment Instrument"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for start, end, mins, kind, text in ASSESS:
    cells = tbl.add_row().cells
    set_cell(cells[0], f"{start}–{end}", bold=True, size=9.5, fill=ASSESS_FILL)
    set_cell(cells[1], f"{mins} min", size=9.5, fill=ASSESS_FILL)
    set_cell(cells[2], text, size=9.5, fill=ASSESS_FILL)
fix_widths(tbl, [0.95, 0.7, 5.15])

# ---- Activity reference
H("Activity Reference (aligned to the learning outcomes)", 1)
tt = doc.add_table(rows=0, cols=4); tt.style = "Table Grid"
hdr = tt.add_row().cells
for i, h in enumerate(["Topic", "Weighting", "LO", "Activities"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for tp in C.TOPICS:
    a = [x for x in ACT if x["topic"] == tp["num"]]
    c = tt.add_row().cells
    set_cell(c[0], f"Topic {tp['code']}: {tp['title']}", bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(c[1], tp["weighting"], size=9.5, fill=TOPIC_FILL)
    set_cell(c[2], a[0]["lo"] if a else "", size=9.5, fill=TOPIC_FILL)
    set_cell(c[3], ", ".join(str(x["num"]) for x in a), size=9.5)
    c[0].width = Inches(2.9); c[1].width = Inches(0.9); c[2].width = Inches(0.6); c[3].width = Inches(2.4)

doc.add_paragraph()
tt = doc.add_table(rows=0, cols=3); tt.style = "Table Grid"
hdr = tt.add_row().cells
for i, h in enumerate(["#", "Activity", "Learning outcome and TSC reference"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for a in ACT:
    c = tt.add_row().cells
    set_cell(c[0], str(a["num"]), bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(c[1], a["title"], size=9.5)
    set_cell(c[2], f"{a['lo']} — {a['objective']}", size=9)
    c[0].width = Inches(0.45); c[1].width = Inches(2.6); c[2].width = Inches(3.75)

# ---- Resources
H("Training Resources", 1)
for r in [f"Trainer slide deck (PPTX and PDF) — {IDX['total']} slides, version {C.VERSION}",
          "Learner Guide (DOCX and PDF) — full step-by-step procedures for all 20 activities",
          "This Lesson Plan (DOCX and PDF)",
          "InDesign lab files, distributed through the LMS at https://lms-tms.tertiaryinfotech.com",
          "One workstation per learner with Adobe InDesign 2024 or later, signed in to Creative Cloud",
          "Projector or large display for trainer demonstration; whiteboard for the client-brief discussion"]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(r).font.size = Pt(10.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT = os.path.join(REPO, "courseware", f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved", OUT)
